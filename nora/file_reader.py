"""
File reader – supports Excel, CSV, PDF, Word and plain text.

Returns a unified FileContent object with extracted text and any
tabular data (as a list of pandas DataFrames).
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pandas as pd

log = logging.getLogger(__name__)


@dataclass
class FileContent:
    """Unified container for content extracted from a file."""

    filename: str
    text: str
    tables: list[pd.DataFrame] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    def summary(self) -> str:
        """Return a short human-readable summary."""
        lines = [f"Fil: {self.filename}"]
        lines.append(f"Tekstlengde: {len(self.text)} tegn")
        if self.tables:
            lines.append(f"Tabeller / ark: {len(self.tables)}")
            for i, df in enumerate(self.tables):
                lines.append(f"  Tabell {i + 1}: {df.shape[0]} rader × {df.shape[1]} kolonner")
        if self.metadata:
            for k, v in self.metadata.items():
                lines.append(f"  {k}: {v}")
        return "\n".join(lines)


def read_file(path: Path) -> FileContent:
    """Read any supported file and return its content."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Finner ikke filen: {path}")

    ext = path.suffix.lower()

    if ext in {".xlsx", ".xls"}:
        return _read_excel(path)
    elif ext == ".csv":
        return _read_csv(path)
    elif ext == ".pdf":
        return _read_pdf(path)
    elif ext in {".docx"}:
        return _read_docx(path)
    elif ext in {".txt", ".md", ".log"}:
        return _read_text(path)
    else:
        log.warning("Ukjent filtype %s – leser som tekst", ext)
        return _read_text(path)


def read_folder(folder: Path, extensions: Optional[set[str]] = None) -> list[FileContent]:
    """Read all supported files in *folder* (non-recursive by default)."""
    folder = Path(folder)
    if not folder.is_dir():
        raise NotADirectoryError(f"Ikke en mappe: {folder}")

    supported = extensions or {".xlsx", ".xls", ".csv", ".pdf", ".docx", ".txt"}
    results: list[FileContent] = []

    for f in sorted(folder.iterdir()):
        if f.is_file() and f.suffix.lower() in supported:
            try:
                log.info("Leser %s", f.name)
                results.append(read_file(f))
            except Exception as exc:
                log.error("Kunne ikke lese %s: %s", f.name, exc)

    return results


# ── Private helpers ────────────────────────────────────────────────────────────


def _read_excel(path: Path) -> FileContent:
    xl = pd.ExcelFile(path)
    tables: list[pd.DataFrame] = []
    text_parts: list[str] = []

    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        tables.append(df)
        text_parts.append(f"[Ark: {sheet}]\n{df.to_string(index=False)}")

    return FileContent(
        filename=path.name,
        text="\n\n".join(text_parts),
        tables=tables,
        metadata={"ark": xl.sheet_names},
    )


def _read_csv(path: Path) -> FileContent:
    # Try UTF-8 first, fall back to latin-1 (common in Norwegian files)
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            df = pd.read_csv(path, encoding=enc, sep=None, engine="python")
            return FileContent(
                filename=path.name,
                text=df.to_string(index=False),
                tables=[df],
                metadata={"encoding": enc},
            )
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Kunne ikke lese CSV-fil: {path}")


def _read_pdf(path: Path) -> FileContent:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError("Installer pdfplumber: pip install pdfplumber") from exc

    text_parts: list[str] = []
    tables: list[pd.DataFrame] = []

    with pdfplumber.open(path) as pdf:
        meta = pdf.metadata or {}
        for i, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text() or ""
            if page_text:
                text_parts.append(f"[Side {i}]\n{page_text}")

            for tbl in page.extract_tables():
                if tbl:
                    df = pd.DataFrame(tbl[1:], columns=tbl[0])
                    tables.append(df)

    return FileContent(
        filename=path.name,
        text="\n\n".join(text_parts),
        tables=tables,
        metadata={k: v for k, v in meta.items() if isinstance(v, (str, int, float))},
    )


def _read_docx(path: Path) -> FileContent:
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise ImportError("Installer python-docx: pip install python-docx") from exc

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables: list[pd.DataFrame] = []

    for tbl in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
            tables.append(df)

    return FileContent(
        filename=path.name,
        text="\n".join(paragraphs),
        tables=tables,
    )


def _read_text(path: Path) -> FileContent:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = path.read_text(encoding=enc)
            return FileContent(filename=path.name, text=text, metadata={"encoding": enc})
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Kunne ikke lese tekstfil: {path}")
